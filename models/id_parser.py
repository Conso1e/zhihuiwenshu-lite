"""
工具函数模块 — 身份证解析、日期处理、文件清理等纯函数。

所有函数均为无副作用的纯函数，不依赖 GUI 状态。
"""

import re
import pandas as pd
from docx import Document


# ====================== 身份证解析 ======================

def parse_id_card(id_number):
    """从 18 位身份证号提取出生年/月/日和性别。
    返回: (year, month, day, gender) 四个字符串，失败时全部为空字符串。
    """
    if not id_number or len(str(id_number)) != 18:
        return '', '', '', ''
    try:
        id_str = str(id_number).strip()
        year = id_str[6:10]
        month = id_str[10:12]
        day = id_str[12:14]
        gender = '男' if int(id_str[16]) % 2 == 1 else '女'
        return year, month, day, gender
    except Exception:
        return '', '', '', ''


# ====================== 文件名清理 ======================

def safe_filename(s):
    """移除文件名中的非法字符。"""
    return re.sub(r'[\\/*?:"<>|]', "_", str(s))


# ====================== 数值转换 ======================

def safe_float(x):
    """安全转换为 float，转换失败返回 0.0。"""
    try:
        return float(x)
    except (ValueError, TypeError):
        return 0.0


# ====================== Word 模板占位符提取 ======================

def extract_placeholders_from_template(template_path):
    """扫描 Word 模板中所有 {{ xxx }} 格式的占位符。
    包括段落、表格单元格、页眉和页脚。
    """
    doc = Document(template_path)
    placeholders = set()
    pattern = r'\{\{\s*(.*?)\s*\}\}'
    for para in doc.paragraphs:
        placeholders.update(re.findall(pattern, para.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    placeholders.update(re.findall(pattern, para.text))
    for section in doc.sections:
        for h in section.header.paragraphs:
            placeholders.update(re.findall(pattern, h.text))
        for f in section.footer.paragraphs:
            placeholders.update(re.findall(pattern, f.text))
    return list(placeholders)


# ====================== 日期处理 ======================

def parse_date_range_to_days(date_str):
    """解析日期区间字符串 (如 '2023-01-01~2024-01-01')，返回天数差。
    支持分隔符: ~ 或 - (需要至少6段以区分单个日期)。
    单个日期或解析失败返回 0。
    """
    if not date_str or pd.isna(date_str):
        return 0

    date_str = str(date_str).strip()
    if not date_str or date_str == '0' or date_str == '':
        return 0

    if hasattr(date_str, 'strftime'):
        return 0

    sep = None
    if '~' in date_str:
        sep = '~'
    elif '-' in date_str:
        parts = date_str.split('-')
        if len(parts) >= 6:
            sep = '-'
        else:
            return 0

    if not sep:
        return 0

    try:
        start_str, end_str = date_str.split(sep, 1)
        start = pd.to_datetime(start_str.strip())
        end = pd.to_datetime(end_str.strip())
        days = (end - start).days
        return max(days, 0)
    except Exception as e:
        print(f"日期解析失败: {date_str}, 错误: {e}")
        return 0


def parse_two_dates_to_days(start_str, end_str):
    """解析两个单独的日期字符串，返回天数差。
    （合并了原代码中两次定义，保留完整的错误处理）
    """
    if not start_str or pd.isna(start_str) or not end_str or pd.isna(end_str):
        return 0
    try:
        if hasattr(start_str, 'strftime'):
            start = start_str
        else:
            start = pd.to_datetime(str(start_str).strip())
        if hasattr(end_str, 'strftime'):
            end = end_str
        else:
            end = pd.to_datetime(str(end_str).strip())
        days = (end - start).days
        return max(days, 0)
    except Exception as e:
        print(f"两个日期解析失败: {start_str} - {end_str}, 错误: {e}")
        return 0


def parse_single_date(date_str):
    """解析单个日期字符串，返回 pandas Timestamp 或 None。"""
    if not date_str or pd.isna(date_str):
        return None
    try:
        if hasattr(date_str, 'strftime'):
            return date_str
        return pd.to_datetime(str(date_str).strip())
    except Exception:
        return None


def extract_start_date(date_str):
    """从日期区间字符串中提取开始日期部分。"""
    if not date_str:
        return ''
    sep = None
    if '-' in str(date_str):
        sep = '-'
    elif '~' in str(date_str):
        sep = '~'
    else:
        return ''
    try:
        return str(date_str).split(sep)[0].strip()
    except Exception:
        return ''


# ====================== 物业费天数获取 ======================

def get_property_days(row):
    """从 DataFrame 行中获取物业服务费欠费天数。
    优先级: 物业服务费欠费周期(日期区间) > 物业费欠费开始/结束日期 > 物业服务费欠费天数(数值)
    """
    # 1. 优先使用合并列 "物业服务费欠费周期"（日期区间）
    prop_date_range = row.get('物业服务费欠费周期', '')
    days = parse_date_range_to_days(prop_date_range) + 1
    if days > 0:
        return days

    # 2. 尝试使用两列 "物业费欠费开始日期" 和 "物业费欠费结束日期"
    start = row.get('物业费欠费开始日期', '')
    end = row.get('物业费欠费结束日期', '')
    days = parse_two_dates_to_days(start, end) + 1
    if days > 0:
        return days

    # 3. 尝试使用数字列 "物业服务费欠费天数"
    try:
        val = row.get('物业服务费欠费天数', 0)
        if val and val != '':
            days = int(float(val))
            if days > 0:
                return days
    except (ValueError, TypeError):
        pass

    return 0


# ====================== 值清理（用于 Word 填充） ======================

def clean_value(val):
    """清理值，去除时间部分，确保纯日期格式 YYYY-MM-DD。
    用于 Word 模板填充前的数据预处理。
    """
    if val is None or pd.isna(val):
        return ''

    if hasattr(val, 'strftime'):
        return val.strftime('%Y-%m-%d')

    if isinstance(val, str):
        val = val.strip()

        if ' ' in val:
            parts = val.split(' ')
            date_part = parts[0]
            if re.match(r'\d{4}-\d{2}-\d{2}', date_part):
                return date_part

        if '.' in val and '-' in val:
            match = re.match(r'(\d{4}-\d{2}-\d{2})', val)
            if match:
                return match.group(1)

        if '/' in val:
            parts = val.split(' ')
            date_part = parts[0].replace('/', '-')
            if re.match(r'\d{4}-\d{2}-\d{2}', date_part):
                return date_part

    return str(val) if val != '' else ''
